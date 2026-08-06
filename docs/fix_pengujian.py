import docx

doc_path = r'c:\laragon\www\garden barbershop\garden_barbershop_finance\docs\Proposal Skripsi v2_updated.docx'
doc = docx.Document(doc_path)

found_pengujian = False
for p in doc.paragraphs:
    text = p.text.strip()
    
    if text == '4.4 Pengujian':
        # Insert introductory text for Pengujian
        p.insert_paragraph_before('4.4 Pengujian').runs[0].bold = True
        p.insert_paragraph_before('Tahapan ini untuk memastikan setiap langkah dalam pengujian sistem berfungsi dengan yang diharapkan, pengujian terhadap sistem ini menggunakan pengujian Black Box Testing dan UAT (User Acceptance Testing). Berikut adalah hasil dari pengujian yang telah dilakukan.')
        p.clear()
        
    elif text == '4.4.1 Black Box Testing':
        p.insert_paragraph_before('4.4.1 Black Box Testing').runs[0].bold = True
        p.insert_paragraph_before('Pengujian ini difokuskan untuk menguji fungsionalitas dan identifikasi masalah setiap fitur yang tidak sesuai dengan hasil yang diharapkan dari berbagai aksi yang dilakukan.')
        p.clear()
        
    elif 'Pengujian sistem menggunakan Black Box Testing. Hasil pengujian dapat dilihat pada Tabel 4.10 berikut.' in text:
        p.clear()
        
    elif 'Tabel 4. 10 Tabel Pengujian Sistem' in text or 'Tabel 4.10 Tabel Pengujian Sistem' in text:
        p.text = 'Tabel 4.10 Pengujian Blackbox Testing'

# To insert UAT, we need to find where Black Box Testing ends. It ends at the Deployment paragraph.
target_deployment = None
for p in doc.paragraphs:
    if 'Sebagai pamungkas dari siklus hidup pengembangan sistem perangkat lunak, sistem dihadapkan pada fase Transition.' in p.text:
        target_deployment = p
        break

if target_deployment:
    # Insert UAT section
    p_uat = target_deployment.insert_paragraph_before('4.4.2 User Acceptance Testing (UAT)')
    p_uat.runs[0].bold = True
    
    target_deployment.insert_paragraph_before('Tahap UAT dilakukan untuk memastikan sistem yang telah dirancang dan dibangun memenuhi kebutuhan pengguna akhir dan memastikan apakah sistem dapat diterima oleh pengguna. Tahapan ini dilakukan dengan pengguna melakukan uji coba sistem secara langsung lalu memberikan penilaian kuesioner, untuk mengetahui tanggapan respon terhadap sistem yang telah dibuat.')
    
    target_deployment.insert_paragraph_before('Tabel 4.11 Kategori Penilaian Bobot')
    # Generate bobot table
    table1 = doc.add_table(rows=1, cols=3)
    table1.style = 'Table Grid'
    h1 = table1.rows[0].cells
    h1[0].text, h1[1].text, h1[2].text = 'Kategori', 'Bobot', 'Keterangan'
    for r in [('Sangat Setuju (SS)', '5', 'Sangat Memuaskan'), ('Setuju (S)', '4', 'Memuaskan'), ('Kurang Setuju (KS)', '3', 'Cukup'), ('Tidak Setuju (TS)', '2', 'Kurang'), ('Sangat Tidak Setuju (STS)', '1', 'Sangat Kurang')]:
        row = table1.add_row().cells
        row[0].text, row[1].text, row[2].text = r
    target_deployment.insert_paragraph_before('')._p.addnext(table1._tbl)
    
    target_deployment.insert_paragraph_before('Tabel 4.12 Pertanyaan Kuesioner')
    # Generate question table
    table2 = doc.add_table(rows=1, cols=2)
    table2.style = 'Table Grid'
    h2 = table2.rows[0].cells
    h2[0].text, h2[1].text = 'No', 'Pertanyaan'
    questions = [
        '1', 'Apakah antarmuka aplikasi mudah dipahami dan digunakan (User Friendly)?',
        '2', 'Apakah fungsionalitas pencatatan pendapatan harian berjalan dengan baik?',
        '3', 'Apakah fitur buku kas umum menyajikan perhitungan saldo yang akurat?',
        '4', 'Apakah rekapitulasi laporan bulanan membantu dalam evaluasi bisnis?',
        '5', 'Apakah aplikasi secara keseluruhan sudah memenuhi kebutuhan Garden Barbershop?'
    ]
    for i in range(0, len(questions), 2):
        row = table2.add_row().cells
        row[0].text, row[1].text = questions[i], questions[i+1]
    target_deployment.insert_paragraph_before('')._p.addnext(table2._tbl)
    
    target_deployment.insert_paragraph_before('Hasil analisis pengujian UAT merupakan hasil perhitungan rata-rata skor dari seluruh pertanyaan berdasarkan responden. Dari hasil perhitungan keseluruhan, dapat disimpulkan bahwa Sistem Informasi Keuangan dan Layanan pada Garden Barbershop ini dapat diterima dengan baik serta sangat membantu dalam proses pengelolaan data barbershop.')

doc.save(doc_path)
print('Pengujian and UAT section updated successfully.')
