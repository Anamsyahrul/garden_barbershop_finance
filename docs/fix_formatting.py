import docx
import shutil
import re

filepath = r'c:\laragon\www\garden barbershop\garden_barbershop_finance\docs\Proposal Skripsi v2.docx'
backup_path = r'c:\laragon\www\garden barbershop\garden_barbershop_finance\docs\Proposal Skripsi v2_backup.docx'

# Restore from backup first to start fresh
shutil.copy(backup_path, filepath)

doc = docx.Document(filepath)

in_bab_4 = False
in_desc = False
in_act = False

desc_counter = 1
act_counter = 1

items_to_match = [
    "Login.",
    "Kelola Akun Pengguna.",
    "Kelola Data Capster.",
    "Kelola Data Layanan.",
    "Input Pendapatan Harian.",
    "Kelola Biaya Operasional.",
    "Kelola Buku Kas Umum.",
    "Lihat Laporan Bulanan."
]

for p in doc.paragraphs:
    text = p.text.strip()
    
    # Track section boundaries
    if "BAB IV" in text:
        in_bab_4 = True
    if "BAB V" in text:
        break
        
    if not in_bab_4:
        continue
        
    # Apply standard heading replacements based on the example structure
    if "HASIL DAN PEMBAHASAN" in text and text == "HASIL DAN PEMBAHASAN":
        p.text = "PEMBAHASAN DAN HASIL"
    elif "Bab ini akan menguraikan secara mendalam mengenai" in text:
        p.text = "Bab ini akan menjelaskan tahapan pengembangan sebuah aplikasi Android pengelolaan pendapatan dan pembagian hasil capster pada Garden Barbershop berbasis Firebase dengan metode pengembangan sistem yang digunakan yaitu metode waterfall. Berikut adalah implementasi Metode waterfall:"
    
    # 4.1 Analisis Kebutuhan
    elif text == "4.1   Fase Inception" or text == "4.1 Fase Inception":
        p.text = "4.1 Analisis Kebutuhan"
    elif text == "4.1.1  Business Modelling Workflow" or text == "4.1.1 Business Modelling Workflow":
        p.text = "4.1.1 Kebutuhan Fungsional"
    elif "Alur kerja business modelling ini bertujuan" in text:
        p.text = "Kebutuhan fungsional merupakan kebutuhan yang menjelaskan proses atau layanan yang harus disediakan oleh sistem agar dapat berjalan sesuai dengan tujuan yang diharapkan. Adapun kebutuhan fungsional pada aplikasi Garden Barbershop Finance adalah sebagai berikut:"
    elif text.startswith("a. Identifikasi Masalah"):
        p.text = "Identifikasi Masalah"
    elif text.startswith("b. Identifikasi Sistem Berjalan"):
        p.text = "Identifikasi Sistem Berjalan"
    
    elif text == "4.1.2  Requirement Workflow" or text == "4.1.2 Requirement Workflow":
        p.text = "4.1.2 Kebutuhan Non-Fungsional"
    elif "Setelah mengidentifikasi masalah dari sistem berjalan" in text:
        p.text = "Kebutuhan Non-Fungsional dilakukan untuk mengetahui spesifikasi kebutuhan untuk sistem. Spesifikasi kebutuhan melibatkan analisis perangkat keras (hardware) dan perangkat lunak (software)."
    elif text.startswith("a. Analisis Kebutuhan Sistem"):
        p.text = "Analisis Kebutuhan Sistem"
    elif text.startswith("b. Fitur Sistem"):
        p.text = "Fitur Sistem"
        
    # 4.2 Desain Sistem
    elif text == "4.2   Fase Elaboration" or text == "4.2 Fase Elaboration":
        p.text = "4.2 Desain Sistem"
    elif text == "4.2.1  Analysis And Design" or text == "4.2.1 Analysis And Design":
        p.text = "Perancangan sistem mencakup arsitektur berbasis framework Flutter dan Firebase, perancangan basis data menggunakan Cloud Firestore, serta perancangan Unified Modeling Language (UML) seperti Use Case Diagram, Activity Diagram, dan Class Diagram untuk menggambarkan alur serta proses sistem yang akan dibangun. Selain itu, perancangan antarmuka menggunakan Flutter untuk menghasilkan tampilan yang responsif pada perangkat Android. Basis data yang dirancang meliputi penyimpanan koleksi dan dokumen semua fitur."
    
    # Restructuring 4.2 sub-items carefully
    elif text == "4.2.1.1 Use Case Diagram":
        p.text = "4.2.1 Use Case Diagram"
    elif text == "4.2.1.2 Deskripsi Use Case":
        p.text = "4.2.2 Deskripsi Use Case"
        in_desc = True
        in_act = False
        desc_counter = 1
    elif text == "4.2.1.3 Activity Diagram":
        p.text = "4.2.3 Activity Diagram"
        in_desc = False
        in_act = True
        act_counter = 1
    elif text == "4.2.1.4 Class Diagram":
        p.text = "4.2.4 Class Diagram"
        in_desc = False
        in_act = False
    elif text == "4.2.1.5 Entity Relationship Diagram":
        p.text = "Entity Relationship Diagram"
    elif text == "4.2.1.6 Skema Basis Data":
        p.text = "Skema Database"
    elif text == "4.2.1.7 Rancangan Tabel":
        p.text = "Rancangan Tabel"
    elif text == "4.2.1.8 Rancangan User Interface":
        p.text = "Rancangan User Interface"
        
    # Formatting list items inside Deskripsi Use Case and Activity Diagram
    elif (in_desc or in_act) and text in items_to_match:
        item_name = text[:-1]  # Remove trailing period
        if in_desc:
            new_text = f"a. {item_name}"
            desc_counter += 1
        else:
            new_text = f"{chr(96 + act_counter)}. {item_name}" # a, b, c, etc.
            act_counter += 1
        
        p.clear()
        run = p.add_run(new_text)
        run.bold = True
        from docx.shared import Inches
        p.paragraph_format.left_indent = Inches(0.25)
        
    # 4.3 Implementasi
    elif text == "4.3 Fase Construction":
        p.text = "4.3 Implementasi"
    elif text == "4.3.1 Implementation Workflow":
        p.text = ""
    elif text == "4.3.1.1 Implementasi Database":
        p.text = "4.3.1 Implementasi Database"
    elif text == "4.3.1.2 Implementasi Antarmuka":
        p.text = "4.3.2 Implementasi Sistem"
        
    # 4.4 Pengujian
    elif text == "4.3.2 Test Workflow":
        p.text = "4.4 Pengujian"
    elif "Pengujian sistem menggunakan Black Box Testing" in text:
        p.insert_paragraph_before("4.4.1 Black Box Testing", p.style)
    elif text == "4.4 Fase Transition":
        p.text = ""
    elif text == "4.4.1 Deployment":
        p.text = "Deployment"

doc.save(filepath)
print("Finished rewriting Proposal Skripsi v2.docx")
