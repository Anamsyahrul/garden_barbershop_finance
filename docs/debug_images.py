import docx
import copy
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc_path = r'c:\laragon\www\garden barbershop\garden_barbershop_finance\docs\Proposal Skripsi v2_updated.docx'
doc = docx.Document(doc_path)

in_bab3 = False
images_to_copy = []

for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if 'BAB III' in text or 'BAB 3' in text:
        in_bab3 = True
    if 'BAB IV' in text or 'BAB 4' in text:
        in_bab3 = False
        
    if in_bab3 and text.startswith('Gambar 3.') and not '\t' in text:
        # Ignore TOC entries that have tabs or numbers at the end
        if len(text.split()) > 2: # At least "Gambar 3.X Title"
            # look for the image in the previous 2 paragraphs
            found_img = None
            for j in range(max(0, i-2), i+1):
                img_p = doc.paragraphs[j]
                if 'w:drawing' in img_p._p.xml:
                    found_img = img_p
                    break
            
            if found_img:
                images_to_copy.append((text, found_img))
                print(f"Collected image for {text}")

print(f"Total collected: {len(images_to_copy)}")
