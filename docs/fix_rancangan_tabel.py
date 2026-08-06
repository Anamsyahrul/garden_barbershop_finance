import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc_path = r'c:\laragon\www\garden barbershop\garden_barbershop_finance\docs\Proposal Skripsi v2_updated.docx'
doc = docx.Document(doc_path)

# State flags
in_db_section = False
paragraphs_to_clear = []
insert_before_p = None

for p in doc.paragraphs:
    text = p.text.strip()
    
    if text == '4.2.6 Skema Basis Data' or text == 'Skema Database' or text == '4.2.6 Skema Database':
        in_db_section = True
        
    if in_db_section:
        if text.startswith('Rancangan User Interface') or text.startswith('4.2.7'):
            in_db_section = False
            insert_before_p = p
        else:
            paragraphs_to_clear.append(p)

# To handle tables that were inserted, clearing paragraphs won't delete the tables themselves!
# We need to delete the tables as well. 
# Finding tables in the document that match our criteria.
# But python-docx doesn't provide an easy way to delete a table directly from the document.
# Instead of dealing with XML deletion, we can just save it to a brand new file by copying elements, 
# or we can clear all content inside the tables and hide borders, but that's messy.
# Wait, if we use python-docx, 	able._element.getparent().remove(table._element) removes a table!

for tbl in doc.tables:
    # If the table has our specific headers, we delete it
    try:
        hdr = tbl.rows[0].cells[0].text
        if hdr == 'Nama Field':
            tbl._element.getparent().remove(tbl._element)
    except:
        pass

for p in paragraphs_to_clear:
    p.clear()

if insert_before_p:
    # Now we insert the new formatted section
    p_title = insert_before_p.insert_paragraph_before('4.2.6 Rancangan Tabel')
    p_title.runs[0].bold = True
    
    tables_data = [
        {
            'num_title': '1. Tabel Users',
            'desc': 'Tabel users digunakan untuk menyimpan data entitas autentikasi pengguna serta pengaturan akses (role).',
            'tbl_title': 'Tabel 4.11 Tabel Users',
            'headers': ['Nama Field', 'Tipe Data', 'Keterangan'],
            'rows': [
                ['id_user', 'String', 'Primary Key, ID unik pengguna'],
                ['username', 'String', 'Nama pengguna untuk login'],
                ['password', 'String', 'Kata sandi pengguna'],
                ['name', 'String', 'Nama lengkap pengguna'],
                ['role', 'String', 'Peran pengguna (Admin/Pemilik/Capster)'],
                ['status', 'String', 'Status keaktifan akun'],
                ['id_capster', 'String', 'Foreign Key, ID capster yang terhubung dengan akun ini (jika ada)']
            ]
        },
        {
            'num_title': '2. Tabel Capsters',
            'desc': 'Tabel capsters digunakan untuk menyimpan data identitas pekerja cukur (capster) yang beroperasi.',
            'tbl_title': 'Tabel 4.12 Tabel Capsters',
            'headers': ['Nama Field', 'Tipe Data', 'Keterangan'],
            'rows': [
                ['id_capster', 'String', 'Primary Key, ID unik capster'],
                ['nama_capster', 'String', 'Nama lengkap capster'],
                ['no_hp', 'String', 'Nomor handphone capster'],
                ['status', 'String', 'Status keaktifan capster']
            ]
        },
        {
            'num_title': '3. Tabel Layanan',
            'desc': 'Tabel layanan digunakan untuk mengelola data jenis jasa pangkas rambut beserta harga yang ditawarkan.',
            'tbl_title': 'Tabel 4.13 Tabel Layanan',
            'headers': ['Nama Field', 'Tipe Data', 'Keterangan'],
            'rows': [
                ['id_layanan', 'String', 'Primary Key, ID unik layanan'],
                ['kode_layanan', 'String', 'Kode layanan'],
                ['nama_layanan', 'String', 'Nama layanan'],
                ['harga', 'Integer', 'Harga layanan'],
                ['kategori', 'String', 'Kategori layanan'],
                ['status', 'String', 'Status layanan (aktif/tidak)']
            ]
        },
        {
            'num_title': '4. Tabel Buku Kas',
            'desc': 'Tabel buku kas digunakan untuk mencatat dan mengelola riwayat seluruh transaksi penerimaan dan pengeluaran kas.',
            'tbl_title': 'Tabel 4.14 Tabel Buku Kas',
            'headers': ['Nama Field', 'Tipe Data', 'Keterangan'],
            'rows': [
                ['id_kas', 'String', 'Primary Key, ID unik pencatatan kas'],
                ['id_user', 'String', 'Foreign Key, ID pengguna yang mencatat'],
                ['tanggal', 'String', 'Tanggal transaksi'],
                ['uraian', 'String', 'Uraian atau deskripsi transaksi'],
                ['akun', 'String', 'Kategori akun transaksi'],
                ['penerimaan', 'Integer', 'Nominal pemasukan kas'],
                ['pengeluaran', 'Integer', 'Nominal pengeluaran kas'],
                ['saldo', 'Integer', 'Saldo akhir kas'],
                ['keterangan', 'String', 'Keterangan tambahan']
            ]
        },
        {
            'num_title': '5. Tabel Operasional',
            'desc': 'Tabel operasional digunakan untuk mencatat setiap biaya operasional rutin seperti tagihan listrik, sewa, dan pengeluaran lainnya.',
            'tbl_title': 'Tabel 4.15 Tabel Operasional',
            'headers': ['Nama Field', 'Tipe Data', 'Keterangan'],
            'rows': [
                ['id_operasional', 'String', 'Primary Key, ID unik biaya operasional'],
                ['id_user', 'String', 'Foreign Key, ID pengguna yang mencatat'],
                ['bulan', 'String', 'Bulan pengeluaran operasional'],
                ['nama_biaya', 'String', 'Nama biaya operasional'],
                ['akun', 'String', 'Kategori akun biaya'],
                ['nominal', 'Integer', 'Jumlah pengeluaran'],
                ['keterangan', 'String', 'Keterangan tambahan']
            ]
        },
        {
            'num_title': '6. Tabel Pendapatan Harian',
            'desc': 'Tabel pendapatan harian digunakan untuk mencatat setoran harian dari masing-masing capster setelah jadwal kerjanya selesai.',
            'tbl_title': 'Tabel 4.16 Tabel Pendapatan Harian',
            'headers': ['Nama Field', 'Tipe Data', 'Keterangan'],
            'rows': [
                ['id_pendapatan', 'String', 'Primary Key, ID unik pendapatan'],
                ['id_user', 'String', 'Foreign Key, ID pengguna yang mencatat'],
                ['id_capster', 'String', 'Foreign Key, ID capster yang bertugas'],
                ['tanggal', 'String', 'Tanggal pencatatan'],
                ['nama_capster', 'String', 'Nama capster'],
                ['pendapatan', 'Integer', 'Jumlah nominal pendapatan tunai'],
                ['cs', 'Integer', 'Jumlah customer shaving (cs)'],
                ['cu', 'Integer', 'Jumlah customer cut (cu)'],
                ['total_customer', 'Integer', 'Total customer yang dilayani']
            ]
        }
    ]
    
    last_element = p_title._p
    
    for tbl_data in tables_data:
        p_num = insert_before_p.insert_paragraph_before(tbl_data['num_title'])
        # To match the requested style, we probably don't need bold but we can bold the title
        # p_num.runs[0].bold = True
        last_element.addnext(p_num._p)
        last_element = p_num._p
        
        p_desc = insert_before_p.insert_paragraph_before(tbl_data['desc'])
        last_element.addnext(p_desc._p)
        last_element = p_desc._p
        
        p_tbl_title = insert_before_p.insert_paragraph_before(tbl_data['tbl_title'])
        p_tbl_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        last_element.addnext(p_tbl_title._p)
        last_element = p_tbl_title._p
        
        table = doc.add_table(rows=1, cols=len(tbl_data['headers']))
        table.style = 'Table Grid'
        
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(tbl_data['headers']):
            hdr_cells[i].text = h
            hdr_cells[i].paragraphs[0].runs[0].font.bold = True
            
        for r_data in tbl_data['rows']:
            row_cells = table.add_row().cells
            for i, val in enumerate(r_data):
                row_cells[i].text = val
                
        last_element.addnext(table._tbl)
        last_element = table._tbl

doc.save(doc_path)
print('Rancangan Tabel replaced successfully.')
