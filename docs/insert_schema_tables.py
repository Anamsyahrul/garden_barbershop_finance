import docx

doc_path = r'c:\laragon\www\garden barbershop\garden_barbershop_finance\docs\Proposal Skripsi v2.docx'
doc = docx.Document(doc_path)

tables_data = [
    {
        "title": "Tabel 4.11 Rancangan Tabel users",
        "headers": ["Nama Field", "Tipe Data", "Keterangan"],
        "rows": [
            ["id_user", "String", "Primary Key, ID unik pengguna"],
            ["username", "String", "Nama pengguna untuk login"],
            ["password", "String", "Kata sandi pengguna"],
            ["name", "String", "Nama lengkap pengguna"],
            ["role", "String", "Peran pengguna (Admin/Pemilik/Capster)"],
            ["status", "String", "Status keaktifan akun"],
            ["id_capster", "String", "Foreign Key, ID capster yang terhubung dengan akun ini (jika ada)"]
        ]
    },
    {
        "title": "Tabel 4.12 Rancangan Tabel capsters",
        "headers": ["Nama Field", "Tipe Data", "Keterangan"],
        "rows": [
            ["id_capster", "String", "Primary Key, ID unik capster"],
            ["nama_capster", "String", "Nama lengkap capster"],
            ["no_hp", "String", "Nomor handphone capster"],
            ["status", "String", "Status keaktifan capster"]
        ]
    },
    {
        "title": "Tabel 4.13 Rancangan Tabel layanan",
        "headers": ["Nama Field", "Tipe Data", "Keterangan"],
        "rows": [
            ["id_layanan", "String", "Primary Key, ID unik layanan"],
            ["kode_layanan", "String", "Kode layanan"],
            ["nama_layanan", "String", "Nama layanan"],
            ["harga", "Integer", "Harga layanan"],
            ["kategori", "String", "Kategori layanan"],
            ["status", "String", "Status layanan (aktif/tidak)"]
        ]
    },
    {
        "title": "Tabel 4.14 Rancangan Tabel buku_kas",
        "headers": ["Nama Field", "Tipe Data", "Keterangan"],
        "rows": [
            ["id_kas", "String", "Primary Key, ID unik pencatatan kas"],
            ["id_user", "String", "Foreign Key, ID pengguna yang mencatat"],
            ["tanggal", "String", "Tanggal transaksi"],
            ["uraian", "String", "Uraian atau deskripsi transaksi"],
            ["akun", "String", "Kategori akun transaksi"],
            ["penerimaan", "Integer", "Nominal pemasukan kas"],
            ["pengeluaran", "Integer", "Nominal pengeluaran kas"],
            ["saldo", "Integer", "Saldo akhir kas"],
            ["keterangan", "String", "Keterangan tambahan"]
        ]
    },
    {
        "title": "Tabel 4.15 Rancangan Tabel operasional",
        "headers": ["Nama Field", "Tipe Data", "Keterangan"],
        "rows": [
            ["id_operasional", "String", "Primary Key, ID unik biaya operasional"],
            ["id_user", "String", "Foreign Key, ID pengguna yang mencatat"],
            ["bulan", "String", "Bulan pengeluaran operasional"],
            ["nama_biaya", "String", "Nama biaya operasional"],
            ["akun", "String", "Kategori akun biaya"],
            ["nominal", "Integer", "Jumlah pengeluaran"],
            ["keterangan", "String", "Keterangan tambahan"]
        ]
    },
    {
        "title": "Tabel 4.16 Rancangan Tabel pendapatan_harian",
        "headers": ["Nama Field", "Tipe Data", "Keterangan"],
        "rows": [
            ["id_pendapatan", "String", "Primary Key, ID unik pendapatan"],
            ["id_user", "String", "Foreign Key, ID pengguna yang mencatat"],
            ["id_capster", "String", "Foreign Key, ID capster yang bertugas"],
            ["tanggal", "String", "Tanggal pencatatan"],
            ["nama_capster", "String", "Nama capster"],
            ["pendapatan", "Integer", "Jumlah nominal pendapatan tunai"],
            ["cs", "Integer", "Jumlah customer shaving (cs)"],
            ["cu", "Integer", "Jumlah customer cut (cu)"],
            ["total_customer", "Integer", "Total customer yang dilayani"]
        ]
    }
]

# Delete existing paragraph entries that were just list points
for p in doc.paragraphs:
    text = p.text.strip()
    if text.startswith("1. Collection 'users'") or \
       text.startswith("2. Collection 'capsters'") or \
       text.startswith("3. Collection 'layanan'") or \
       text.startswith("4. Collection 'buku_kas'") or \
       text.startswith("5. Collection 'operasional'") or \
       text.startswith("6. Collection 'pendapatan_harian'"):
        p.clear()

target_p = None
for p in doc.paragraphs:
    if "Secara spesifik, rancangan struktur dokumen NoSQL di dalam Firebase dijabarkan secara terperinci sebagai berikut" in p.text:
        target_p = p
        break

if target_p:
    last_element = target_p._p
    
    for tbl_data in tables_data:
        new_p = doc.add_paragraph(tbl_data["title"])
        new_p.style = 'Normal'
        # To make it bold and centered maybe? 
        # For simplicity, we just use the default paragraph style
        last_element.addnext(new_p._p)
        last_element = new_p._p
        
        table = doc.add_table(rows=1, cols=len(tbl_data["headers"]))
        table.style = 'Table Grid'
        
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(tbl_data["headers"]):
            hdr_cells[i].text = h
            # Optional: bold headers
            hdr_cells[i].paragraphs[0].runs[0].font.bold = True
            
        for r_data in tbl_data["rows"]:
            row_cells = table.add_row().cells
            for i, val in enumerate(r_data):
                row_cells[i].text = val
                
        last_element.addnext(table._tbl)
        last_element = table._tbl

doc.save(doc_path)
print("Tables created successfully.")
