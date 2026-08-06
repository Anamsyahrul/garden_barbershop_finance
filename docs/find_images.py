import docx

doc_path = r'c:\laragon\www\garden barbershop\garden_barbershop_finance\docs\Proposal Skripsi v2_updated.docx'
doc = docx.Document(doc_path)

for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if 'Gambar 3.3 Halaman Login' in text or 'Gambar 3.4 Dashboard Admin' in text:
        print(f"Found text at para {i}: '{text}'")
        # Check surrounding paragraphs for drawings
        for j in range(max(0, i-2), min(len(doc.paragraphs), i+3)):
            p_xml = doc.paragraphs[j]._p.xml
            if 'w:drawing' in p_xml:
                print(f"  -> Paragraph {j} contains an image!")
