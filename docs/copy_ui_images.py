import docx
import copy
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc_path = r'c:\laragon\www\garden barbershop\garden_barbershop_finance\docs\Proposal Skripsi v2_updated.docx'
doc = docx.Document(doc_path)

# 1. Collect images from Bab 3
# Titles to look for in Bab 3
bab3_titles = [
    "Gambar 3.3 Halaman Login",
    "Gambar 3.4 Dashboard Admin",
    "Gambar 3.5 Input Pendapatan Harian",
    "Gambar 3.6 Buku Kas Umum",
    "Gambar 3.7 Input Operasional",
    "Gambar 3.8 Laporan Pembagian Hasil",
    "Gambar 3.9 Kelola Akun Pengguna"
]

images_to_copy = []
for title in bab3_titles:
    for i, p in enumerate(doc.paragraphs):
        # ensure we match the exact caption, ignoring TOC
        if title in p.text and len(p.text.strip()) == len(title):
            # The image should be at i-1
            img_p = doc.paragraphs[i-1]
            if 'w:drawing' in img_p._p.xml:
                images_to_copy.append((title, img_p))
                break

# Check if we got all images
print(f"Collected {len(images_to_copy)} images from Bab 3.")

# 2. Find Rancangan User Interface section in Bab 4
in_rui = False
paragraphs_to_clear = []
insert_before_p = None

for p in doc.paragraphs:
    text = p.text.strip()
    if text == "Rancangan User Interface" or text == "Rancangan User Inferface":
        in_rui = True
        
    if in_rui:
        if text.startswith("4.3 Implementasi"):
            in_rui = False
            insert_before_p = p
        else:
            paragraphs_to_clear.append(p)

for p in paragraphs_to_clear:
    p.clear()

if insert_before_p:
    # Insert heading
    p_heading = insert_before_p.insert_paragraph_before("Rancangan User Interface")
    p_heading.runs[0].bold = True
    
    # Insert intro text
    insert_before_p.insert_paragraph_before("Rancangan desain sistem ini disusun sebagai acuan dalam proses pengembangan agar sistem yang dihasilkan sesuai dengan tujuan, mudah digunakan, serta mampu mendukung proses pengelolaan data keuangan secara efektif dan efisien. Namun, rancangan ini bersifat fleksibel dan dapat mengalami perubahan atau penyesuaian seiring dengan perkembangan kebutuhan serta kondisi pada tahap implementasi sistem.")
    
    last_element = doc.paragraphs[-1]._p # temporary
    # Actually, let's keep track of the element to addnext
    # Wait, insert_paragraph_before creates a new paragraph BEFORE the target.
    # We can just insert a blank paragraph, then replace its xml with the image xml!
    # That is safer.
    
    start_num = 17
    for title, img_p in images_to_copy:
        # Title will be Gambar 4.x [Name]
        name = title.replace("Gambar 3.", "").split(" ", 1)[1] # "Halaman Login"
        new_title = f"Gambar 4.{start_num} {name}"
        
        # Insert image
        new_img_p = insert_before_p.insert_paragraph_before("")
        new_img_p._p.getparent().replace(new_img_p._p, copy.deepcopy(img_p._p))
        
        # Insert caption
        new_caption_p = insert_before_p.insert_paragraph_before(new_title)
        new_caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        start_num += 1

doc.save(doc_path)
print("Rancangan User Interface updated successfully.")
