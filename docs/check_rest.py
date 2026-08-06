import docx
doc_path = r'c:\laragon\www\garden barbershop\garden_barbershop_finance\docs\Proposal Skripsi v2_updated.docx'
doc = docx.Document(doc_path)
for i, p in enumerate(doc.paragraphs):
    if '4.3 Implementasi' in p.text:
        for j in range(i, min(i+50, len(doc.paragraphs))):
            print(doc.paragraphs[j].text.strip())
        break
