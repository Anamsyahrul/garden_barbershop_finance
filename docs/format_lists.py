import docx

filepath = r'c:\laragon\www\garden barbershop\garden_barbershop_finance\docs\Proposal Skripsi v2.docx'
doc = docx.Document(filepath)

in_desc = False
in_act = False

desc_counter = 1
act_counter = 1

items_to_match = [
    "Login.",
    "Kelola Akun Pengguna.",
    "Kelola Data Capster.",
    "Kelola Data Layanan.",
    "Input Pendapatan Harian.",
    "Kelola Biaya Operasional.",
    "Kelola Buku Kas Umum.",
    "Lihat Laporan Bulanan."
]

for p in doc.paragraphs:
    text = p.text.strip()
    
    if text == "Deskripsi Use Case":
        in_desc = True
        in_act = False
        desc_counter = 1
        continue
    elif text == "4.2.2 Activity Diagram":
        in_desc = False
        in_act = True
        act_counter = 1
        continue
    elif text == "4.2.3 Class Diagram" or text == "Entity Relationship Diagram":
        in_desc = False
        in_act = False
        continue

    if in_desc or in_act:
        if text in items_to_match:
            # Reformat the paragraph
            item_name = text[:-1]  # Remove trailing period
            if in_desc:
                new_text = f"a. {item_name}"
                desc_counter += 1
            else:
                new_text = f"{chr(96 + act_counter)}. {item_name}" # a, b, c, etc.
                act_counter += 1
            
            # Clear existing runs and add the new bold text
            p.clear()
            run = p.add_run(new_text)
            run.bold = True
            
            # Add some indentation if possible (just using python-docx indent)
            from docx.shared import Inches
            p.paragraph_format.left_indent = Inches(0.25)

doc.save(filepath)
print("Finished reformatting Use Case and Activity numbering.")
