import docx
import copy
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc_path = r'c:\laragon\www\garden barbershop\garden_barbershop_finance\docs\Proposal Skripsi v2_updated.docx'
doc = docx.Document(doc_path)

# Extract images from Bab 3
in_bab3 = False
ui_images = []

for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if 'BAB III' in text or 'BAB 3' in text:
        in_bab3 = True
    if 'BAB IV' in text or 'BAB 4' in text:
        in_bab3 = False
        
    if in_bab3 and text.startswith('Gambar 3.') and not '\t' in text:
        if 'Tahapan' not in text and 'Tempat' not in text:
            found_img = None
            for j in range(max(0, i-2), i+1):
                img_p = doc.paragraphs[j]
                if 'w:drawing' in img_p._p.xml:
                    found_img = img_p
                    break
            
            if found_img:
                # clean the title to just get the name
                # e.g., "Gambar 3.4 Dashboard Admin" -> "Dashboard Admin"
                # e.g., "Gambar 3. Input Pendapatan" -> "Input Pendapatan"
                name_parts = text.split(' ', 1)[1]
                if name_parts.startswith('3.') or name_parts.startswith('.'):
                    name_parts = name_parts.split(' ', 1)[-1]
                ui_images.append((name_parts.strip(), found_img))

print(f"Collected {len(ui_images)} UI images from Bab 3.")

# Rewrite Rancangan User Interface
in_rui = False
paragraphs_to_clear = []
insert_before_p = None

for p in doc.paragraphs:
    text = p.text.strip()
    if text == "Rancangan User Interface" or text == "Rancangan User Inferface":
        in_rui = True
        
    if in_rui:
        if text.startswith("4.3 Implementasi"):
            in_rui = False
            insert_before_p = p
        else:
            paragraphs_to_clear.append(p)

for p in paragraphs_to_clear:
    p.clear()

if insert_before_p:
    p_heading = insert_before_p.insert_paragraph_before("Rancangan User Interface")
    p_heading.runs[0].bold = True
    
    insert_before_p.insert_paragraph_before("Rancangan desain sistem ini disusun sebagai acuan dalam proses pengembangan agar sistem yang dihasilkan sesuai dengan tujuan, mudah digunakan, serta mampu mendukung proses pengelolaan data keuangan secara efektif dan efisien. Namun, rancangan ini bersifat fleksibel dan dapat mengalami perubahan atau penyesuaian seiring dengan perkembangan kebutuhan serta kondisi pada tahap implementasi sistem.")
    
    start_num = 17
    for name, img_p in ui_images:
        new_title = f"Gambar 4.{start_num} {name}"
        
        # We need to insert the title FIRST or IMAGE FIRST?
        # In 'contoh penulisan baru', it looks like:
        # Gambar 4.14 Halaman Login
        # [Image]
        # Wait, usually the caption for a Figure is BELOW the image in APA/IEEE, but sometimes above.
        # Looking at 'contoh_penulisan.txt', the caption was below or above?
        # "Gambar 4. 14 Halaman Login
        # Gambar 4. 15 Halaman Dashboard" 
        # Actually, in the output, it's listed sequentially. Let's just put the image, then the title below it.
        
        # Insert image clone
        new_img_p = insert_before_p.insert_paragraph_before("")
        new_img_p._p.getparent().replace(new_img_p._p, copy.deepcopy(img_p._p))
        
        # Insert caption
        new_caption_p = insert_before_p.insert_paragraph_before(new_title)
        new_caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        start_num += 1

doc.save(doc_path)
print("Rancangan User Interface fully updated with images!")
