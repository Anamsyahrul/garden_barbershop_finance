import docx

doc_path = r'c:\laragon\www\garden barbershop\garden_barbershop_finance\docs\Proposal Skripsi v2.docx'
doc = docx.Document(doc_path)

table_data = [
    ['No', 'Skenario Pengujian', 'Hasil yang Diharapkan', 'Kesimpulan'],
    ['1', 'Login', '', ''],
    ['', 'Memasukkan username dan password yang benar.', 'Sistem menerima akses dan mengarahkan ke halaman Dashboard sesuai role.', '[ ✔ ] Valid'],
    ['', 'Memasukkan username atau password yang salah.', 'Sistem menolak akses dan menampilkan pesan Username atau password salah.', '[ ✔ ] Valid'],
    ['', 'Mengosongkan form username dan password, lalu klik Login.', 'Sistem menampilkan pesan validasi form tidak boleh kosong.', '[ ✔ ] Valid'],
    ['2', 'Kelola Akun Pengguna', '', ''],
    ['', 'Menambahkan akun pengguna baru dengan data yang lengkap.', 'Data akun baru tersimpan dan muncul pada daftar pengguna.', '[ ✔ ] Valid'],
    ['', 'Mengubah/memperbarui data akun pengguna.', 'Sistem menyimpan perubahan dan memperbarui data akun.', '[ ✔ ] Valid'],
    ['', 'Menghapus salah satu data akun pengguna.', 'Akun terhapus dari daftar dan tidak dapat login kembali.', '[ ✔ ] Valid'],
    ['3', 'Kelola Data Capster', '', ''],
    ['', 'Menambahkan data capster baru.', 'Data capster ditambahkan dan tampil pada daftar capster.', '[ ✔ ] Valid'],
    ['', 'Menonaktifkan status capster.', 'Status capster menjadi tidak aktif dan tidak dapat dipilih saat input harian.', '[ ✔ ] Valid'],
    ['4', 'Kelola Data Layanan', '', ''],
    ['', 'Menambah jenis layanan baru beserta harga.', 'Layanan baru tersimpan dan tampil pada daftar layanan.', '[ ✔ ] Valid'],
    ['', 'Mengubah nominal harga pada layanan yang sudah ada.', 'Harga layanan tersebut berhasil diperbarui di database.', '[ ✔ ] Valid'],
    ['5', 'Input Pendapatan Harian', '', ''],
    ['', 'Menyimpan data setoran harian capster.', 'Rincian pendapatan tersimpan dan tercatat di Buku Kas.', '[ ✔ ] Valid'],
    ['', 'Menyimpan setoran dengan nilai Rp 0 atau form kosong.', 'Sistem menolak proses dan memunculkan peringatan.', '[ ✔ ] Valid'],
    ['6', 'Kelola Biaya Operasional', '', ''],
    ['', 'Mencatat pengeluaran operasional.', 'Biaya operasional tersimpan dan tercatat sebagai pengeluaran di kas.', '[ ✔ ] Valid'],
    ['7', 'Kelola Buku Kas Umum', '', ''],
    ['', 'Melihat mutasi penerimaan dan pengeluaran.', 'Sistem menampilkan riwayat transaksi (Buku Kas) dengan akurat.', '[ ✔ ] Valid'],
    ['', 'Memastikan kalkulasi nilai saldo akhir.', 'Sistem menghitung selisih akhir saldo kas secara otomatis tanpa kesalahan.', '[ ✔ ] Valid'],
    ['8', 'Lihat Laporan Bulanan', '', ''],
    ['', 'Mem-filter laporan keuangan berdasarkan bulan tertentu.', 'Sistem menyajikan rekap pendapatan, pengeluaran, dan laba bersih dengan benar.', '[ ✔ ] Valid']
]

target_p = None
for p in doc.paragraphs:
    if 'Pengujian sistem menggunakan Black Box Testing' in p.text or 'Tabel 4.10 Tabel Pengujian Sistem' in p.text:
        target_p = p
        
if target_p is None:
    print('Could not find the target paragraph.')
else:
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'No'
    hdr_cells[1].text = 'Skenario Pengujian'
    hdr_cells[2].text = 'Hasil yang Diharapkan'
    hdr_cells[3].text = 'Kesimpulan'
    
    for row_data in table_data[1:]:
        row_cells = table.add_row().cells
        for i, text in enumerate(row_data):
            row_cells[i].text = text

    target_p._p.addnext(table._tbl)
    doc.save(doc_path)
    print('Table successfully inserted and saved.')
