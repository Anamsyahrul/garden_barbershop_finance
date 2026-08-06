import docx

doc_path = r'c:\laragon\www\garden barbershop\garden_barbershop_finance\docs\Proposal Skripsi v2_updated.docx'
doc = docx.Document(doc_path)

# Update paragraphs
for p in doc.paragraphs:
    if p.text:
        # Paragraphs text replacements
        new_text = p.text
        # Capsters -> Capster
        new_text = new_text.replace('Tabel Capsters', 'Tabel Capster')
        new_text = new_text.replace('Tabel capsters', 'Tabel capster')
        new_text = new_text.replace('Database Capsters', 'Database Capster')
        
        # buku_kas -> buku_kas_umum
        new_text = new_text.replace('Tabel buku_kas', 'Tabel buku_kas_umum')
        new_text = new_text.replace('Tabel Buku Kas', 'Tabel Buku Kas Umum')
        new_text = new_text.replace('Database Buku Kas', 'Database Buku Kas Umum')
        new_text = new_text.replace('Tabel buku kas', 'Tabel buku kas umum')
        
        if p.text != new_text:
            # Reconstruct the paragraph keeping basic formatting by replacing the whole text
            # if it's a simple text string. If we do p.text = new_text we lose inline formatting,
            # but for our generated headings and simple sentences it's fine.
            # But wait, we can also iterate runs to be safe
            for run in p.runs:
                if 'Capsters' in run.text or 'capsters' in run.text or 'Buku Kas' in run.text or 'buku_kas' in run.text or 'buku kas' in run.text:
                    rt = run.text
                    rt = rt.replace('Capsters', 'Capster')
                    rt = rt.replace('capsters', 'capster')
                    rt = rt.replace('Buku Kas', 'Buku Kas Umum')
                    rt = rt.replace('buku_kas', 'buku_kas_umum')
                    rt = rt.replace('buku kas', 'buku kas umum')
                    run.text = rt

# Update tables if any contain these texts (for instance table headers or cell text)
for tbl in doc.tables:
    for row in tbl.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    if 'Capsters' in run.text or 'capsters' in run.text or 'Buku Kas' in run.text or 'buku_kas' in run.text or 'buku kas' in run.text:
                        rt = run.text
                        rt = rt.replace('Capsters', 'Capster')
                        rt = rt.replace('capsters', 'capster')
                        rt = rt.replace('Buku Kas', 'Buku Kas Umum')
                        rt = rt.replace('buku_kas', 'buku_kas_umum')
                        rt = rt.replace('buku kas', 'buku kas umum')
                        run.text = rt

doc.save(doc_path)
print("Docx updated successfully.")
