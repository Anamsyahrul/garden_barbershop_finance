import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc_path = r'c:\laragon\www\garden barbershop\garden_barbershop_finance\docs\Proposal Skripsi v2_updated.docx'
doc = docx.Document(doc_path)

# We want to insert BAB V right before DAFTAR PUSTAKA
# Find the DAFTAR PUSTAKA paragraph
target_p = None
for p in doc.paragraphs:
    if 'DAFTAR PUSTAKA' in p.text.upper():
        target_p = p
        break

if target_p:
    # Insert before target_p
    # docx doesn't have a direct insert_paragraph_before on the Document object,
    # but we can use target_p.insert_paragraph_before(text)
    
    # We want a page break before BAB V
    # target_p.insert_paragraph_before().add_run().add_break(docx.enum.text.WD_BREAK.PAGE)
    # Actually, let's just insert the paragraphs in reverse order or sequentially
    
    p_bab5 = target_p.insert_paragraph_before('BAB V\nKESIMPULAN DAN SARAN')
    p_bab5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Make it bold if we can, but styles are easier. We'll just set it bold
    for run in p_bab5.runs:
        run.bold = True
        
    p_kesimpulan_title = target_p.insert_paragraph_before('5.1 Kesimpulan')
    p_kesimpulan_title.runs[0].bold = True
    
    p_kesimpulan_text = target_p.insert_paragraph_before('Berdasarkan hasil analisis, perancangan, dan implementasi yang telah dilakukan pada pengembangan Aplikasi Keuangan dan Layanan Garden Barbershop Berbasis Android, maka dapat ditarik beberapa kesimpulan sebagai berikut:\n'
    '1. Sistem ini berhasil mempermudah pencatatan pendapatan harian, pengeluaran operasional, dan pengelolaan buku kas umum yang sebelumnya masih rentan terhadap kesalahan manusia (human error) menjadi lebih terintegrasi dan akurat.\n'
    '2. Penggunaan arsitektur NoSQL menggunakan Firebase Firestore memberikan efisiensi yang tinggi dalam sinkronisasi data secara real-time antar pengguna (Admin, Pemilik, dan Capster).\n'
    '3. Berdasarkan hasil pengujian Black Box, seluruh fungsionalitas utama aplikasi berjalan dengan valid dan dapat menghasilkan laporan bulanan yang transparan untuk memudahkan evaluasi kinerja bisnis maupun evaluasi bagi masing-masing Capster.')
    p_kesimpulan_text.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    p_saran_title = target_p.insert_paragraph_before('5.2 Saran')
    p_saran_title.runs[0].bold = True
    
    p_saran_text = target_p.insert_paragraph_before('Dalam rangka pengembangan sistem lebih lanjut di masa mendatang, penulis memberikan beberapa saran yang dapat dipertimbangkan, yaitu:\n'
    '1. Mengembangkan fitur pemesanan layanan (booking/reservasi) secara daring (online) yang terintegrasi langsung dengan pelanggan (customer).\n'
    '2. Menambahkan fitur notifikasi (push notification) kepada pengguna terkait ringkasan laporan harian maupun jadwal/shift kerja.\n'
    '3. Mengintegrasikan sistem pembayaran dengan pihak ketiga (Payment Gateway) agar Garden Barbershop dapat melayani transaksi pembayaran non-tunai (cashless) secara otomatis dari aplikasi.')
    p_saran_text.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Add a page break before DAFTAR PUSTAKA to keep it on a new page (Optional but recommended)
    # target_p.runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)

    doc.save(doc_path)
    print('BAB V inserted successfully.')
else:
    print('Could not find DAFTAR PUSTAKA')
