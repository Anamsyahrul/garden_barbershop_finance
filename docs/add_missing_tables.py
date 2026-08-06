import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc_path = r'c:\laragon\www\garden barbershop\garden_barbershop_finance\docs\Proposal Skripsi v2_updated.docx'
doc = docx.Document(doc_path)

tables_data = [
    {
        'num_title': '7. Tabel Laporan Bulanan',
        'desc': 'Tabel laporan bulanan adalah tabel agregasi yang digunakan untuk menyimpan rekapitulasi total pendapatan bersih dan pengeluaran setiap bulan, sehingga sistem tidak perlu membaca data harian dari awal saat menampilkan grafik.',
        'tbl_title': 'Tabel 4.17 Tabel Laporan Bulanan',
        'headers': ['Nama Field', 'Tipe Data', 'Keterangan'],
        'rows': [
            ['id_laporan', 'String', 'Primary Key, ID unik laporan'],
            ['bulan', 'String', 'Bulan rekapitulasi (contoh: Januari)'],
            ['tahun', 'String', 'Tahun rekapitulasi (contoh: 2026)'],
            ['total_pendapatan', 'Integer', 'Total seluruh pendapatan di bulan tersebut'],
            ['total_pengeluaran', 'Integer', 'Total seluruh operasional di bulan tersebut'],
            ['laba_bersih', 'Integer', 'Pendapatan dikurangi pengeluaran']
        ]
    },
    {
        'num_title': '8. Tabel Rekap Customer',
        'desc': 'Tabel rekap customer digunakan sebagai tabel agregasi untuk meringkas jumlah pelanggan yang dilayani oleh setiap capster per periode (hari/bulan), mempercepat proses pembuatan laporan kinerja.',
        'tbl_title': 'Tabel 4.18 Tabel Rekap Customer',
        'headers': ['Nama Field', 'Tipe Data', 'Keterangan'],
        'rows': [
            ['id_rekap', 'String', 'Primary Key, ID unik rekapitulasi'],
            ['id_capster', 'String', 'Foreign Key, ID capster yang dievaluasi'],
            ['periode', 'String', 'Periode rekapan (contoh: Juni 2026)'],
            ['total_cs', 'Integer', 'Total layanan Customer Shaving (cs)'],
            ['total_cu', 'Integer', 'Total layanan Customer Cut (cu)'],
            ['total_keseluruhan', 'Integer', 'Total keseluruhan pelanggan (cs + cu)']
        ]
    }
]

# Find where to insert (after Table 4.16)
target_p = None
for p in doc.paragraphs:
    if 'Tabel 4.16 Tabel Pendapatan Harian' in p.text:
        target_p = p
        break

if target_p:
    # We have to skip past the actual table element to insert after it
    # We find the table that immediately follows target_p
    current_element = target_p._p.getnext() # This should be the w:tbl
    
    for tbl_data in tables_data:
        # Create paragraphs and tables directly on the document, then move their XML 
        # to the correct position (after current_element)
        p_num = doc.add_paragraph(tbl_data['num_title'])
        current_element.addnext(p_num._p)
        current_element = p_num._p
        
        p_desc = doc.add_paragraph(tbl_data['desc'])
        current_element.addnext(p_desc._p)
        current_element = p_desc._p
        
        p_tbl_title = doc.add_paragraph(tbl_data['tbl_title'])
        p_tbl_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        current_element.addnext(p_tbl_title._p)
        current_element = p_tbl_title._p
        
        table = doc.add_table(rows=1, cols=len(tbl_data['headers']))
        table.style = 'Table Grid'
        
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(tbl_data['headers']):
            hdr_cells[i].text = h
            hdr_cells[i].paragraphs[0].runs[0].font.bold = True
            
        for r_data in tbl_data['rows']:
            row_cells = table.add_row().cells
            for i, val in enumerate(r_data):
                row_cells[i].text = val
                
        current_element.addnext(table._tbl)
        current_element = table._tbl

doc.save(doc_path)
print("Added missing tables to document successfully.")
