import docx
from docx.shared import Pt
import os

path = r'c:\laragon\www\garden barbershop\garden_barbershop_finance\docs\Proposal Skripsi v2_final.docx'
doc = docx.Document(path)

# Data from our Markdown artifact
data = [
    ["No", "Pertanyaan", "SS (5)", "S (4)", "KS (3)", "TS (2)", "STS (1)", "Total Skor"],
    ["1", "Apakah antarmuka aplikasi mudah dipahami dan digunakan (User Friendly)?", "3", "2", "0", "0", "0", "23"],
    ["2", "Apakah fungsionalitas pencatatan pendapatan harian berjalan dengan baik?", "4", "1", "0", "0", "0", "24"],
    ["3", "Apakah fitur buku kas umum umum menyajikan perhitungan saldo yang akurat?", "3", "2", "0", "0", "0", "23"],
    ["4", "Apakah rekapitulasi laporan bulanan membantu dalam evaluasi bisnis?", "5", "0", "0", "0", "0", "25"],
    ["5", "Apakah aplikasi secara keseluruhan sudah memenuhi kebutuhan Garden Barbershop?", "4", "1", "0", "0", "0", "24"],
    ["", "TOTAL SKOR KESELURUHAN", "", "", "", "", "", "119"]
]

# We need to append the table at the end of the UAT section.
# The UAT section is at the very end of chapter 4, before chapter 5.
# Let's just append it to the very end of the document, or better, 
# find the paragraph "Hasil analisis pengujian UAT merupakan hasil perhitungan..."
target_idx = -1
for i, p in enumerate(doc.paragraphs):
    if "Hasil analisis pengujian UAT merupakan hasil perhitungan rata-rata skor" in p.text:
        target_idx = i
        break

if target_idx != -1:
    # Insert before the conclusion paragraph
    p = doc.paragraphs[target_idx].insert_paragraph_before("Tabel 4.13 Hasil Kuesioner UAT (Simulasi 5 Responden)")
    table = doc.paragraphs[target_idx].insert_paragraph_before("").insert_paragraph_before("") # Hack to insert table before a paragraph isn't direct in python-docx, it's easier to just add at the end or use private methods.

# Let's use the safer append method at the end of the document, but we want it in Chapter 4.
# Actually python-docx can add a table at the end of the document.
# Since Chapter 5 is at the end, let's find "BAB V" and insert before it.
p_bab5 = None
for p in doc.paragraphs:
    if "BAB V" in p.text:
        p_bab5 = p
        break

# Function to insert table before a paragraph
def insert_table_before(paragraph, data):
    # Create an empty table at the end of the document
    tbl = doc.add_table(rows=len(data), cols=len(data[0]))
    tbl.style = 'Table Grid'
    for r_idx, row in enumerate(data):
        for c_idx, cell_val in enumerate(row):
            tbl.cell(r_idx, c_idx).text = cell_val
    
    # Move the table before the paragraph
    p_parent = paragraph._element.getparent()
    tbl_element = tbl._element
    p_parent.insert(p_parent.index(paragraph._element), tbl_element)
    
    # Return the created table
    return tbl

if p_bab5:
    p_title = p_bab5.insert_paragraph_before("Tabel 4.13 Hasil Simulasi Kuesioner UAT (5 Responden)")
    insert_table_before(p_bab5, data)
    
    calc_text = (
        "\nPerhitungan Persentase Kelayakan UAT:\n"
        "- Skor Maksimal (Ideal) = Jumlah Responden (5) × Jumlah Pertanyaan (5) × Skor Tertinggi (5) = 125\n"
        "- Total Skor Didapat = 119\n"
        "- Persentase = (119 / 125) × 100% = 95.2%\n\n"
        "Berdasarkan hasil pengujian simulasi kepada 5 responden, didapatkan persentase sebesar 95.2% yang masuk dalam kategori Sangat Layak / Sangat Setuju."
    )
    p_bab5.insert_paragraph_before(calc_text)
    
doc.save(path)
print("UAT data successfully written to the docx file.")
