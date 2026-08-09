import docx
import re
import shutil

path = r'c:\laragon\www\garden barbershop\garden_barbershop_finance\docs\Proposal Skripsi v2_final.docx'
backup_path = r'c:\laragon\www\garden barbershop\garden_barbershop_finance\docs\Proposal Skripsi v2_final_BACKUP.docx'
shutil.copy2(path, backup_path)

doc = docx.Document(path)

foreign_words = [
    'Database', 'User Interface', 'Dashboard', 'Login', 'Username', 'Password', 'Role', 
    'Real-time', 'Activity Diagram', 'Class Diagram', 'Entity Relationship Diagram', 
    'Black Box Testing', 'User Acceptance Testing', 'Collections', 'Documents', 
    'NoSQL', 'SQL', 'UAT', 'Customer', 'Role-based access', 'Logout', 'Test Case', 
    'User Friendly', 'Framework', 'Web Console', 'Firestore', 'MySQL', 'JSON', 'real-time',
    'Black Box', 'White Box', 'Deployment', 'Build process', 'Android Application Package', 'apk',
    'Firebase', 'Console'
]

foreign_words.sort(key=len, reverse=True)
pattern = re.compile(r'\b(?:' + '|'.join(map(re.escape, foreign_words)) + r')\b', re.IGNORECASE)

def copy_format(source_run, target_run, force_italic=False):
    target_run.bold = source_run.bold
    target_run.italic = True if force_italic else source_run.italic
    target_run.underline = source_run.underline
    target_run.font.name = source_run.font.name
    target_run.font.size = source_run.font.size
    if source_run.font.color and source_run.font.color.rgb:
        target_run.font.color.rgb = source_run.font.color.rgb

def process_paragraphs(paragraphs):
    for p in paragraphs:
        runs = list(p.runs)
        for run in runs:
            if run.italic:
                continue
            text = run.text
            if not text:
                continue
                
            matches = list(pattern.finditer(text))
            if not matches:
                continue
                
            run.text = ""
            current_idx = 0
            current_xml_element = run._r
            
            for match in matches:
                start, end = match.span()
                if start > current_idx:
                    new_run = p.add_run(text[current_idx:start])
                    copy_format(run, new_run, force_italic=False)
                    current_xml_element.addnext(new_run._r)
                    current_xml_element = new_run._r
                    
                new_run_match = p.add_run(text[start:end])
                copy_format(run, new_run_match, force_italic=True)
                current_xml_element.addnext(new_run_match._r)
                current_xml_element = new_run_match._r
                
                current_idx = end
                
            if current_idx < len(text):
                new_run_after = p.add_run(text[current_idx:])
                copy_format(run, new_run_after, force_italic=False)
                current_xml_element.addnext(new_run_after._r)
                current_xml_element = new_run_after._r

process_paragraphs(doc.paragraphs)

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            process_paragraphs(cell.paragraphs)

doc.save(path)
print("Sukses memiringkan kata!")
