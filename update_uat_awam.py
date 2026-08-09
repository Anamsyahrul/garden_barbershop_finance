import docx

path = r'c:\laragon\www\garden barbershop\garden_barbershop_finance\docs\Proposal Skripsi v2_final.docx'
doc = docx.Document(path)

new_questions = [
    "Apakah desain tampilan (antarmuka) aplikasi ini menarik dan nyaman dilihat?",
    "Apakah menu dan tombol navigasi dalam aplikasi ini mudah dipahami dan diakses?",
    "Apakah teks dan informasi yang disajikan di dalam aplikasi dapat dibaca dengan jelas?",
    "Apakah aplikasi ini terlihat profesional untuk digunakan sebagai alat bantu keuangan?",
    "Secara keseluruhan, apakah aplikasi ini mudah digunakan meskipun oleh orang awam?"
]

# Update Table 27 (index 26) - Pertanyaan
table_27 = doc.tables[26]
for i, q in enumerate(new_questions):
    table_27.cell(i+1, 1).text = q

# Update Table 28 (index 27) - Simulasi Hasil
table_28 = doc.tables[27]
for i, q in enumerate(new_questions):
    table_28.cell(i+1, 1).text = q

doc.save(path)
print("Pertanyaan UAT di Word berhasil diganti untuk orang awam.")
